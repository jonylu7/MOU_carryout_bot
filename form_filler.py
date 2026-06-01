"""
carry_out_automation/form_filler.py

將 CarryOutRequest 列表填入攜出申請單 .docx 範本，
並插入電子簽章圖片，最後另存新檔。

依賴：
    pip install python-docx Pillow google-genai
"""

from __future__ import annotations

import io
import logging
import re
from datetime import datetime
from pathlib import Path
from typing import Optional

from google import genai
from docx import Document
from docx.shared import Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import config
from sheets_reader import (
    CarryOutRequest,
    read_sheet_values,
    export_sheet_range_as_png,
    export_workbook_as_pngs,
)

logger = logging.getLogger(__name__)

_CN_NUMS        = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]
_EXAMPLE_KWDS   = ("工作表", "圖檔", "圖片")


# ─────────────────────────────────────────────────
# 輔助：日期解析 & 西元年 → 民國年字串
# ─────────────────────────────────────────────────

_DATE_FMTS = ("%Y-%m-%d %H:%M:%S", "%Y/%m/%d", "%Y-%m-%d", "%m/%d/%Y", "%Y年%m月%d日")


def _parse_date(date_str: str) -> Optional[datetime]:
    """嘗試多種格式解析日期字串，回傳 datetime 或 None。"""
    for fmt in _DATE_FMTS:
        try:
            return datetime.strptime(str(date_str).strip(), fmt)
        except ValueError:
            continue
    return None


def to_roc_date(date_str: str) -> str:
    """
    輸入 '2026/05/29' 或 '2026-06-05 00:00:00' → '115 年 05 月 29 日'
    若解析失敗則原樣回傳。
    """
    dt = _parse_date(date_str)
    if dt is None:
        return date_str
    roc_year = dt.year - config.ROC_YEAR_OFFSET
    return f"{roc_year} 年 {dt.month:02d} 月 {dt.day:02d} 日"


def to_roc_today() -> str:
    today = datetime.today()
    roc_year = today.year - config.ROC_YEAR_OFFSET
    return f"{roc_year} 年 {today.month:02d} 月 {today.day:02d} 日"


# ─────────────────────────────────────────────────
# Gemini API：整理檔案描述
# ─────────────────────────────────────────────────

def llm_format_file_rows(requests: list[CarryOutRequest]) -> list[dict]:
    """
    用 Gemini 將多筆申請的檔案資訊整理成結構化列表，
    每個元素代表攜出申請單中的一個檔案列：
      { "檔名": ..., "欄項說明": ..., "描述": ... }
    """
    if not config.GEMINI_API_KEY:
        return _simple_format_file_rows(requests)

    rows_text = "\n".join(
        f"- 填寫人：{r.填寫人}，檔名：{r.檔案名稱}，屬性：{r.檔案屬性}，"
        f"說明：{r.格式內容說明}，備註：{r.特殊備註}"
        for r in requests
    )

    prompt = f"""你是一位行政助理，負責整理財政部財政資訊中心的資料攜出申請單。
請將以下各填寫人的攜出檔案資訊，整理成一個 JSON 陣列。
每個元素代表一位申請人（一人一列），若一位填寫人有多個檔案請以換行符 \\n 分隔。
格式如下：
{{
  "檔名": "（填寫人名字縮寫_原始檔名；多個檔案以 \\n 分隔，如：威任_a.xlsx\\n威任_b.csv）",
  "欄項說明": "（每個檔案的類型說明，多個以 \\n 分隔，如：工作表，見範例一\\n數值資料）",
  "描述": "（30字以內，說明攜出目的及使用資料）"
}}

申請資料：
{rows_text}

請直接輸出 JSON 陣列，不要加任何說明文字。"""

    import json
    try:
        client = genai.Client(api_key=config.GEMINI_API_KEY)
        response = client.models.generate_content(
            model=config.GEMINI_MODEL,
            contents=prompt,
        )

        raw = response.text.strip()
        raw = re.sub(r"^```[a-z]*\n?", "", raw)
        raw = re.sub(r"\n?```$", "", raw)

        file_rows = json.loads(raw)
        # 後處理：為工作表/圖檔加範例序號（LLM 可能未加）
        example_map = _build_example_map(requests)
        for row in file_rows:
            fname = row.get("檔名", "")
            for person, num in example_map.items():
                # 以姓名末兩字作為檔名對應依據
                hint = person[-2:] if len(person) >= 2 else person
                if hint in fname:
                    row["欄項說明"] = _with_example(row.get("欄項說明", ""), num)
                    break
        return file_rows
    except Exception as e:
        logger.warning(f"Gemini API 呼叫失敗（{e}），改用純文字整理")
        return _simple_format_file_rows(requests)


def _build_example_map(requests: list[CarryOutRequest]) -> dict[str, str]:
    """
    為有範例截圖需求的填寫人依序分配範例序號（一人一號）。
    條件：範例連結欄位含有效 URL，或 檔案屬性 含工作表/圖檔/圖片關鍵字。
    """
    mapping: dict[str, str] = {}
    counter = 0
    for r in requests:
        if r.填寫人 in mapping:
            continue
        link = str(r.範例連結).strip()
        has_url = link and link.lower() != "nan" and link.startswith("http")
        has_workbook_attr = any(kw in str(r.檔案屬性) for kw in _EXAMPLE_KWDS)
        if has_url or has_workbook_attr:
            if counter < len(_CN_NUMS):
                mapping[r.填寫人] = _CN_NUMS[counter]
                counter += 1
    return mapping   # {person: "一", ...}


def _with_example(attr: str, example_num: str | None) -> str:
    """若屬性為工作表/圖檔且尚無範例標記，附加「，如範例X」。"""
    if example_num and any(kw in attr for kw in _EXAMPLE_KWDS) and "如範例" not in attr:
        return f"{attr}，如範例{example_num}"
    return attr


def _simple_format_file_rows(requests: list[CarryOutRequest]) -> list[dict]:
    """每位申請人的所有檔案合成一列（多檔以換行分隔）；工作表/圖檔自動加範例序號。"""
    example_map = _build_example_map(requests)
    rows = []
    for r in requests:
        files = [f.strip() for f in str(r.檔案名稱).split("\n") if f.strip()]
        attrs = [a.strip() for a in str(r.檔案屬性).split("\n") if a.strip()]
        num   = example_map.get(r.填寫人)
        attr_parts = []
        for i in range(len(files)):
            raw_attr = attrs[i] if i < len(attrs) else (attrs[0] if attrs else "")
            attr_parts.append(_with_example(raw_attr, num))
        rows.append({
            "檔名":     "\n".join(files) if files else "",
            "欄項說明": "\n".join(attr_parts),
            "描述":     r.格式內容說明 or r.特殊備註,
        })
    return rows if rows else [{"檔名": "", "欄項說明": "", "描述": ""}]


# ─────────────────────────────────────────────────
# 範例截圖：下載 Google Sheet → PNG 插入文件
# ─────────────────────────────────────────────────

def _extract_sheet_id(url: str) -> Optional[str]:
    """從 Google Sheets / Drive URL 擷取 file ID。"""
    # spreadsheets URL：https://docs.google.com/spreadsheets/d/{ID}/…
    m = re.search(r"/spreadsheets/d/([a-zA-Z0-9_-]+)", url)
    if m:
        return m.group(1)
    # 一般 Drive URL：https://drive.google.com/file/d/{ID}/…
    m = re.search(r"/d/([a-zA-Z0-9_-]+)", url)
    if m:
        return m.group(1)
    return None


def _smart_crop_ranges_via_gemini(values: list[list]) -> list[str]:
    """
    把試算表 cell 資料（純文字）送 Gemini，請它識別所有獨立的資料表格區域。
    回傳 A1 range 列表（例如 ['A1:E5', 'A8:D12']）。
    失敗時回傳空列表。
    """
    import json as _json

    if not config.GEMINI_API_KEY or not values:
        return []

    # 最多傳前 50 列 × 前 25 欄
    lines = []
    for i, row in enumerate(values[:50]):
        trimmed = [str(c) for c in row[:25]]
        lines.append(f"Row {i + 1}: {trimmed}")
    text_repr = "\n".join(lines)

    prompt = f"""以下是一個 Google 試算表的儲存格內容（格式：Row N: [儲存格1, 儲存格2, ...]）。
這份試算表是填寫人用來示範「攜出資料的欄位格式」，需要截圖放入申請書。

【判斷方式】
試算表中可能有多個子表格，子表格的起點特徵是：
- 該列的儲存格內容是「欄位名稱/描述性標籤」（例如：年份、縣市代碼、所得金額、ID、名稱 等），
  而非數值或 ID 資料
- 這種「標頭列」之前可能有空白列，也可能沒有

請依此規則找出每個子表格的範圍：
1. 每遇到一個「標頭列」（欄位名稱列），就代表一個新子表格的開始
2. 子表格的結束是下一個「標頭列」前一列（或資料結束處）
3. 每個 range 只到最後一個有資料的欄（不要包含右側空白欄）
4. 每個 range 只到最後一個有資料的列（不要包含下方空白列）
5. 若某子表格欄位超過 8 欄，嘗試按邏輯分成左右兩段（各自為一個 range）

回傳純 JSON 陣列，每個元素：{{"range": "A2:E8", "label": "子表格簡短說明"}}
**只回傳 JSON，不要有任何說明文字**。

試算表內容：
{text_repr}"""

    try:
        client = genai.Client(api_key=config.GEMINI_API_KEY)
        resp = client.models.generate_content(model=config.GEMINI_MODEL, contents=prompt)
        raw = resp.text.strip()
        # 去除可能的 markdown code fence
        raw = re.sub(r"^```[a-z]*\n?", "", raw)
        raw = re.sub(r"\n?```$", "", raw)

        data = _json.loads(raw)
        valid_re = re.compile(r"^[A-Z]+\d+:[A-Z]+\d+$", re.IGNORECASE)
        ranges = []
        for item in data:
            if not isinstance(item, dict):
                continue
            rng = item.get("range", "")
            if valid_re.match(rng):
                ranges.append(rng.upper())
                label = item.get("label", "")
                logger.info(f"  識別區域：{rng.upper()}{f'（{label}）' if label else ''}")
        if ranges:
            logger.info(f"Gemini 識別到 {len(ranges)} 個表格區域")
        else:
            logger.warning("Gemini 未能識別任何有效 range，改用全頁 fallback")
        return ranges
    except Exception as e:
        logger.warning(f"Gemini 智慧裁切判斷失敗：{e}，改用全頁 fallback")
        return []


def _autocrop_png(png_bytes: bytes, padding: int = 15) -> bytes:
    """
    用 PIL 去除 PNG 四周的白邊，加回 padding 像素的留白後回傳。
    若 PIL 不可用則原樣回傳。
    """
    try:
        from PIL import Image, ImageChops
        img = Image.open(io.BytesIO(png_bytes)).convert("RGB")
        bg = Image.new("RGB", img.size, (255, 255, 255))
        diff = ImageChops.difference(img, bg)
        bbox = diff.getbbox()
        if bbox:
            x0 = max(0, bbox[0] - padding)
            y0 = max(0, bbox[1] - padding)
            x1 = min(img.width,  bbox[2] + padding)
            y1 = min(img.height, bbox[3] + padding)
            img = img.crop((x0, y0, x1, y1))
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        return buf.getvalue()
    except ImportError:
        logger.debug("PIL 未安裝，跳過自動裁切白邊")
        return png_bytes
    except Exception as e:
        logger.debug(f"PIL 自動裁切失敗：{e}")
        return png_bytes


def _fetch_example_as_pngs(url: str) -> list[bytes]:
    """
    從 Google Sheets / Drive xlsx URL 取得所有工作表截圖，回傳 list[bytes]。
    每個 worksheet 一張圖，空白工作表略過。
    """
    file_id = _extract_sheet_id(url)
    if not file_id:
        logger.warning(f"無法從 URL 擷取 file ID：{url}")
        return []

    try:
        images = export_workbook_as_pngs(file_id)
        images = [_autocrop_png(img) for img in images]
        images = [img for img in images if img]
        if images:
            logger.info(f"✅ 共截取 {len(images)} 張截圖（{file_id}）")
            return images
    except Exception as e:
        logger.warning(f"截圖取得失敗（{file_id}）：{e}")

    return _fetch_fullpage_pngs(file_id)


def _fetch_fullpage_pngs(file_id: str) -> list[bytes]:
    """全頁截圖 fallback。export_sheet_range_as_png 統一處理 Google Sheets 與 xlsx on Drive。"""
    try:
        png = export_sheet_range_as_png(file_id, None)
        png = _autocrop_png(png)
        logger.info(f"✅ 全頁截圖（裁切後）：{len(png)} bytes")
        return [png]
    except Exception as e:
        logger.warning(f"全頁截圖取得失敗（{file_id}）：{e}")
        return []


def _insert_image_to_paragraph(
    para,
    image_bytes: bytes,
    max_width_cm: float = 14.0,
    max_height_cm: float = 7.0,
):
    """
    清空段落後插入圖片，確保不超過最大寬度與最大高度（等比縮放）。
    """
    try:
        from PIL import Image as _PILImage
        img = _PILImage.open(io.BytesIO(image_bytes))
        img_w, img_h = img.size
    except Exception:
        img_w, img_h = 1, 1  # 無法讀取尺寸時 fallback

    max_w = Cm(max_width_cm)
    max_h = Cm(max_height_cm)

    # 計算等比縮放後的尺寸
    if img_w > 0 and img_h > 0:
        ratio = img_h / img_w
        h_if_fit_width = int(max_w * ratio)
        if h_if_fit_width <= max_h:
            width, height = max_w, None   # 以寬度為基準
        else:
            width, height = None, max_h   # 以高度為基準（圖太高）
    else:
        width, height = max_w, None

    for run in para.runs:
        run.text = ""
    run = para.add_run()
    run.add_picture(io.BytesIO(image_bytes), width=width, height=height)


def _add_example_heading(doc: Document, num: str, type_label: str = "工作表"):
    """在文件末尾附加一個「範例N　工作表/圖檔」標題段落。"""
    para = doc.add_paragraph()
    run = para.add_run(f"範例{num}　{type_label}")
    run.bold = True
    # 字型
    run.font.name = "DFKai-SB"
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), "DFKai-SB")


def _fill_example_sections(
    doc: Document,
    requests: list[CarryOutRequest],
    example_map: dict[str, str],
):
    """
    1. 先刪除文件中現有的任何「範例X」段落（模板殘留的或上次填寫的）
    2. 在文件末尾為每個有工作表/圖檔的申請人動態附加：
         - 「範例N　工作表/圖檔」標題段落
         - 每個獨立表格一張截圖段落（Gemini 識別多表格）
    """
    # ── 清除模板/文件中現有 範例X 段落（以防萬一） ──
    found_first = False
    to_delete = []
    for para in doc.paragraphs:
        if not found_first and re.search(r"範例[一二三四五六七八九十]", para.text):
            found_first = True
        if found_first:
            to_delete.append(para)
    for para in to_delete:
        para._p.getparent().remove(para._p)
    if to_delete:
        logger.info(f"已清除既有 範例 區段（{len(to_delete)} 個段落）")

    if not example_map:
        return

    # ── 建立映射 ──────────────────────────────────
    num_to_person: dict[str, str] = {v: k for k, v in example_map.items()}
    person_to_req: dict[str, CarryOutRequest] = {r.填寫人: r for r in requests}
    person_to_link: dict[str, str] = {}
    for r in requests:
        link = str(r.範例連結).strip()
        if link and link.lower() != "nan" and link.startswith("http"):
            person_to_link[r.填寫人] = link

    # ── 依序附加各範例 ─────────────────────────────
    first = True
    for num in _CN_NUMS:
        person = num_to_person.get(num)
        if not person:
            break   # example_map 是連續分配的，後面也不會有了

        if first:
            from docx.enum.text import WD_BREAK
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            first = False

        req = person_to_req.get(person)
        attr = str(req.檔案屬性).strip() if req else ""
        type_label = "圖檔" if any(kw in attr for kw in ("圖檔", "圖片")) else "工作表"

        _add_example_heading(doc, num, type_label)

        link = person_to_link.get(person)
        if not link:
            logger.info(f"範例{num}（{person}）無範例連結，略過截圖")
            continue

        images = _fetch_example_as_pngs(link)
        if not images:
            note = doc.add_paragraph()
            note.add_run("（截圖取得失敗，請手動補充）")
            logger.warning(f"範例{num}（{person}）截圖取得失敗")
            continue

        for img_bytes in images:
            _insert_image_to_paragraph(doc.add_paragraph(), img_bytes)

        logger.info(f"✅ 範例{num}（{person}）插入 {len(images)} 張截圖")


# ─────────────────────────────────────────────────
# 主要：填寫 .docx 範本
# ─────────────────────────────────────────────────

def _select_template(topic: str) -> Path:
    for key, fname in config.TEMPLATES.items():
        if key in topic:
            return config.TEMPLATE_DIR / fname
    return config.TEMPLATE_DIR / config.DEFAULT_TEMPLATE


def _collapse_para(para, new_text: str):
    """將段落所有 run 的文字合成一個：run[0] = new_text，其餘清空。"""
    if not para.runs:
        para.add_run(new_text)
        return
    para.runs[0].text = new_text
    for run in para.runs[1:]:
        run.text = ""


def _replace_in_paragraph(para, old: str, new: str):
    """跨 run 的精確字串替換（保留格式）。"""
    full = "".join(r.text for r in para.runs)
    if old not in full:
        return
    _collapse_para(para, full.replace(old, new))


def _replace_regex_in_paragraph(para, pattern: str, replacement: str):
    """跨 run 的 regex 替換（保留格式）。"""
    full = "".join(r.text for r in para.runs)
    new_full = re.sub(pattern, replacement, full)
    if new_full != full:
        _collapse_para(para, new_full)


def fill_form(
    requests: list[CarryOutRequest],
    carry_date: str,
    carry_time: str,
    topic: str,
    output_path: Path,
) -> Path:
    """
    填寫攜出申請單，儲存到 output_path，回傳 output_path。
    """
    template_path = _select_template(topic)
    if not template_path.exists():
        raise FileNotFoundError(f"找不到範本：{template_path}")

    doc = Document(str(template_path))

    # 申請人 & 攜出人：固定為協調人（盧威任）
    applicant     = config.DEFAULT_APPLICANT
    apply_date_roc = to_roc_today()
    carry_date_roc = to_roc_date(carry_date)

    # ── 段落替換 ──────────────────────────────────
    # ROC 日期樣式：三位數字 年 … 月 … 日（runs 可能拆得很碎，用 regex）
    roc_date_pattern = r"\d{3}\s*年\s*\d+\s*月\s*\d+\s*日"
    for para in doc.paragraphs:
        _replace_in_paragraph(para, "張逸芩", applicant)   # 申請人 & 指定攜出人
        _replace_regex_in_paragraph(para, roc_date_pattern, apply_date_roc)
        _replace_regex_in_paragraph(para, r"YYYYMMDD", datetime.today().strftime("%Y%m%d"))

    # ── 填寫 Table 0（主資料表）──────────────────
    if doc.tables:
        _fill_main_table(doc.tables[0], carry_date_roc, carry_time, requests, topic)

    # ── 插入範例截圖 ──────────────────────────────
    _fill_example_sections(doc, requests, _build_example_map(requests))

    # ── 插入電子簽章 ─────────────────────────────
    if config.SIGNATURE_IMAGE_PATH.exists():
        _insert_signature(doc, applicant)
    else:
        logger.warning(f"找不到簽章圖片：{config.SIGNATURE_IMAGE_PATH}，略過簽章")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    logger.info(f"攜出申請單已儲存：{output_path}")
    return output_path


def _fill_main_table(table, carry_date_roc: str, carry_time: str,
                     requests: list[CarryOutRequest], topic: str):
    """填寫 Table 0。"""
    file_rows   = llm_format_file_rows(requests)
    total_files = sum(
        max(1, len([f for f in row.get("檔名", "").split("\n") if f.strip()]))
        for row in file_rows
    )

    # 攜出時間格式化
    time_str = carry_time.strip()
    t_obj = None
    for time_fmt in ("%H:%M:%S", "%H:%M"):
        try:
            t_obj = datetime.strptime(time_str, time_fmt)
            break
        except ValueError:
            continue
    time_formatted = f"{t_obj.hour} 時 {t_obj.minute:02d} 分" if t_obj else time_str

    # 研究題目
    topic_desc_map = {
        "子題一":      "調降營所稅及促產條例落日之政策效果",
        "子題一子議題": "稅制改革引起之地下經濟",
        "子題二":      "台灣租稅與年金改革及世代間不平等",
        "稅收估計":    "使用數據分析進行各稅別稅收估計",
    }
    research_topic = next(
        (v for k, v in topic_desc_map.items() if k in topic), topic
    )

    for row in table.rows:
        # 去重（合併儲存格在 python-docx 中會重複出現）
        seen = set()
        unique_cells = []
        for cell in row.cells:
            if id(cell) not in seen:
                seen.add(id(cell))
                unique_cells.append(cell)

        for cell in unique_cells:
            t = cell.text

            # 預定攜出時間：collapse 所有 run，直接寫入新文字
            if "預定攜出時間" in t:
                new_text = f"預定攜出時間： {carry_date_roc} {time_formatted}"
                for para in cell.paragraphs:
                    if "預定攜出時間" in "".join(r.text for r in para.runs):
                        _collapse_para(para, new_text)
                        break

            # 件數
            if "檔案資料件(筆)數" in t:
                for para in cell.paragraphs:
                    _replace_regex_in_paragraph(para, r"\d+件", f"{total_files}件")

            # 研究題目
            if "所需資料" in t:
                for para in cell.paragraphs:
                    _replace_regex_in_paragraph(
                        para, r"「.+?」", f"「{research_topic}」"
                    )

    _fill_file_rows(table, file_rows)


def _fill_file_rows(table, file_rows: list[dict]):
    """清除舊資料列，再依序填入新檔案列。列數不足時自動在頁腳前插入新列。"""
    from copy import deepcopy

    # 找標題列，保存頁腳列的 _tr 元素（不用 index，因插入後 index 會變）
    header_row_idx = None
    footer_tr = None
    for i, row in enumerate(table.rows):
        row_text = " ".join(c.text for c in row.cells)
        if ("檔案(file)" in row_text or "file" in row_text.lower()) and header_row_idx is None:
            header_row_idx = i
        elif header_row_idx is not None and ("所需資料" in row_text or "件(筆)數" in row_text):
            footer_tr = row._tr
            break

    if header_row_idx is None:
        logger.warning("找不到標題列，跳過檔案列填寫")
        return

    data_start = header_row_idx + 2   # 跳過標題列 + 欄位子標題列

    def _data_rows():
        rows = table.rows
        end = next((i for i, r in enumerate(rows) if r._tr is footer_tr), len(rows))
        return list(rows[data_start:end])

    # 先清空所有舊資料列
    for row in _data_rows():
        seen = set()
        for cell in row.cells:
            if id(cell) not in seen:
                seen.add(id(cell))
                _set_cell_text(cell, "")

    # 保存最後一個資料列的 XML 作為新列的複製來源
    initial = _data_rows()
    template_tr = initial[-1]._tr if initial else None

    for i, frow in enumerate(file_rows):
        rows = _data_rows()
        if i >= len(rows):
            if template_tr is None:
                logger.warning(f"無法新增列（無範本列），略過第 {i+1} 筆")
                continue
            new_tr = deepcopy(template_tr)
            if footer_tr is not None:
                footer_tr.addprevious(new_tr)
            else:
                table._tbl.append(new_tr)
            rows = _data_rows()

        seen = set()
        unique = []
        for c in rows[i].cells:
            if id(c) not in seen:
                seen.add(id(c))
                unique.append(c)
        if len(unique) >= 2:
            _set_cell_text(unique[0], frow.get("檔名", ""),      font_name="Times New Roman")
            _set_cell_text(unique[1], frow.get("欄項說明", ""),   font_name="DFKai-SB")
            _set_cell_text(unique[-1], frow.get("描述", ""),      font_name="DFKai-SB")

    logger.info(f"已填入 {len(file_rows)} 個檔案列")


def _set_cell_text(cell, text: str, font_name: str = "DFKai-SB"):
    """清空 cell 並設定文字與字體（保留第一個 paragraph 的格式）。"""
    if not cell.paragraphs:
        return
    para = cell.paragraphs[0]
    if not para.runs:
        run = para.add_run(text)
    else:
        para.runs[0].text = text
        for run in para.runs[1:]:
            run.text = ""
        run = para.runs[0]
    # 設定字體
    run.font.name = font_name
    # 同時設定東亞字型（Word 對中文字型的實際控制點）
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"),    font_name)
    rFonts.set(qn("w:hAnsi"),    font_name)


# ─────────────────────────────────────────────────
# 電子簽章插入
# ─────────────────────────────────────────────────

def _insert_signature(doc: Document, applicant: str):
    """在含有申請人姓名的段落後方插入簽章圖片。"""
    for para in doc.paragraphs:
        if applicant in para.text and "申請人" in para.text:
            run = para.add_run()
            run.add_picture(
                str(config.SIGNATURE_IMAGE_PATH),
                width=Cm(config.SIGNATURE_WIDTH_CM),
            )
            logger.info(f"已在「{applicant}」後方插入電子簽章")
            return
    logger.warning("找不到申請人段落，簽章插入略過")


# ─────────────────────────────────────────────────
# docx → PDF 轉換
# ─────────────────────────────────────────────────

def docx_to_pdf(docx_path: Path) -> Path:
    """將 docx 轉換為 PDF，回傳 PDF Path。依賴 docx2pdf（需系統有 Word 或 LibreOffice）。"""
    from docx2pdf import convert
    pdf_path = docx_path.with_suffix(".pdf")
    convert(str(docx_path), str(pdf_path))
    logger.info(f"已轉換為 PDF：{pdf_path}")
    return pdf_path


# ─────────────────────────────────────────────────
# 建立攜出資料夾
# ─────────────────────────────────────────────────

def create_carryout_folder(carry_date: str, topic: str) -> Path:
    """
    建立 攜出資料夾(統一放這)/YYYYMM/YYYYMMDD_topic/ 資料夾，回傳 Path。
    """
    dt = _parse_date(carry_date) or datetime.today()
    ym_tag     = dt.strftime("%Y%m")
    date_tag   = dt.strftime("%Y%m%d")
    folder_path = config.CARRYOUT_FOLDER / ym_tag / f"{date_tag}{topic}攜出"
    folder_path.mkdir(parents=True, exist_ok=True)
    logger.info(f"已建立攜出資料夾：{folder_path}")
    return folder_path
